import os
import re
import io
import zipfile
from datetime import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxEditor:
    """
    Advanced DOCX Surgical Engine
    Supports Tracked Changes, Comments, and Structural XML Editing.
    """
    
    NAMESPACE = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }

    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = Document(file_path)
        with open(file_path, 'rb') as f:
            self.bytes = f.read()

    # --- READING CAPABILITIES ---
    
    def read_text(self):
        """Returns all text in the document."""
        return "\n".join([p.text for p in self.doc.paragraphs])

    def read_paragraphs(self):
        """Returns a list of paragraphs with metadata."""
        return [
            {"index": i, "style": p.style.name, "text": p.text}
            for i, p in enumerate(self.doc.paragraphs)
        ]

    def read_tables(self):
        """Returns all table data as 2D arrays."""
        tables_data = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables_data.append(table_data)
        return tables_data

    # --- TRACKED CHANGES (SURGICAL XML) ---

    def _get_xml(self, part_name='word/document.xml'):
        with zipfile.ZipFile(io.BytesIO(self.bytes)) as z:
            return z.read(part_name).decode('utf-8')

    def replace_tracked(self, old_text, new_text, author="Kali AI"):
        """Performs a search-and-replace with tracked changes (redline)."""
        date_str = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        
        # Load the raw XML
        xml_content = self._get_xml()
        
        # Escape characters for XML
        old_esc = old_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        new_esc = new_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Pattern for redlining
        replacement = (
            f'<w:del w:id="0" w:author="{author}" w:date="{date_str}">'
            f'<w:r><w:delText>{old_esc}</w:delText></w:r></w:del>'
            f'<w:ins w:id="1" w:author="{author}" w:date="{date_str}">'
            f'<w:r><w:t>{new_esc}</w:t></w:r></w:ins>'
        )
        
        if old_esc in xml_content:
            xml_content = xml_content.replace(old_esc, replacement)
            self._update_xml(xml_content)
            return True
        return False

    def insert_tracked(self, target_text, insert_text, position="after", author="Kali AI"):
        """Inserts text with tracking relative to a target phrase."""
        date_str = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        xml_content = self._get_xml()
        
        target_esc = target_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        insert_esc = insert_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        tracking = f'<w:ins w:id="2" w:author="{author}" w:date="{date_str}"><w:r><w:t>{insert_esc}</w:t></w:r></w:ins>'
        
        if position == "after":
            new_val = target_esc + tracking
        else:
            new_val = tracking + target_esc
            
        if target_esc in xml_content:
            xml_content = xml_content.replace(target_esc, new_val)
            self._update_xml(xml_content)
            return True
        return False

    def add_comment(self, comment_text, on_text, author="Kali AI"):
        """Adds a comment to a specific phrase (simplified logic)."""
        # Note: True OOXML comments require entries in comments.xml and rels.
        # This implementation adds it as a visible inline Note for compatibility.
        for para in self.doc.paragraphs:
            if on_text in para.text:
                run = para.add_run(f" [COMMENT by {author}: {comment_text}]")
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.italic = True
                return True
        return False

    # --- FORMATTING & STRUCTURE ---

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def add_bullet(self, text):
        self.doc.add_paragraph(text, style='List Bullet')

    def add_table(self, headers, rows):
        table = self.doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)

    def format_paragraph(self, index, bold=False, italic=False, color_hex=None):
        if index < len(self.doc.paragraphs):
            para = self.doc.paragraphs[index]
            for run in para.runs:
                run.bold = bold
                run.italic = italic
                if color_hex:
                    run.font.color.rgb = RGBColor.from_string(color_hex)

    # --- UTILS ---

    def _update_xml(self, new_xml, part_name='word/document.xml'):
        """Repacks the DOCX with modified XML."""
        in_buf = io.BytesIO(self.bytes)
        out_buf = io.BytesIO()
        
        with zipfile.ZipFile(in_buf, 'r') as zin:
            with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == part_name:
                        zout.writestr(item, new_xml)
                    else:
                        zout.writestr(item, zin.read(item.filename))
        
        self.bytes = out_buf.getvalue()
        # Reload python-docx object from new bytes
        self.doc = Document(io.BytesIO(self.bytes))

    def save(self, output_path):
        """Saves the final surgical result."""
        # Use the underlying bytes if surgery was performed, otherwise the doc object
        with open(output_path, 'wb') as f:
            f.write(self.bytes)

def ai_edit_document(file_path, instructions):
    """
    Entry point for AI agents to edit documents.
    instructions: List of dicts [{"action": "replace", "old": "...", "new": "..."}, ...]
    """
    editor = DocxEditor(file_path)
    for task in instructions:
        action = task.get("action")
        if action == "replace_tracked":
            editor.replace_tracked(task["old"], task["new"])
        elif action == "insert_tracked":
            editor.insert_tracked(task["target"], task["text"], task.get("position", "after"))
        elif action == "comment":
            editor.add_comment(task["text"], task["on"])
        elif action == "heading":
            editor.add_heading(task["text"], task.get("level", 1))
        elif action == "bullet":
            editor.add_bullet(task["text"])
            
    output_path = f"revised_{os.path.basename(file_path)}"
    editor.save(output_path)
    return output_path
