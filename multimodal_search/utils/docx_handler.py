from docx import Document
import os
import shutil
import io
import re
import zipfile
import tempfile
from datetime import datetime

def load_doc(source):
    # Handle both paths (strings) and streams (BytesIO)
    if isinstance(source, str) and not os.path.exists(source):
        raise FileNotFoundError(f"File not found: {source}")
    return Document(source)

def verify_docx_bytes(data):
    """Verifies that a binary docx stream is architecturally sound."""
    try:
        source = io.BytesIO(data)
        Document(source)
        return True
    except Exception:
        return False

def save_doc(doc, path):
    # Backup before overwrite
    backup = None
    if os.path.exists(path):
        timestamp = datetime.now().strftime('%H%M%S')
        backup = path.replace(".docx", f"_backup_{timestamp}.docx")
        shutil.copy2(path, backup)
    
    try:
        doc.save(path)
        # Verify immediately
        Document(path)
    except Exception as e:
        # Restore backup if corrupt or failed
        if backup and os.path.exists(backup):
            shutil.copy2(backup, path)
            raise RuntimeError(f"Save produced corrupt file — backup restored: {str(e)}")
        else:
            raise RuntimeError(f"Critical Save Failure: {str(e)}")

def edit_docx_and_return(template_path, replacements, output_path):
    """
    Classic Find-and-Replace and Structural Edit
    """
    doc = Document(template_path)
    
    # 1. Body Text
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
                
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)
                            
    doc.save(output_path)
    return output_path
