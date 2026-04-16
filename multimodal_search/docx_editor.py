"""
docx_editor.py  —  Drop this file into your project.
Gives your AI the ability to READ, EDIT, and CREATE .docx files
just like Claude does — tracked changes, comments, formatting, tables, and more.

QUICK START:
    from docx_editor import DocxEditor
    editor = DocxEditor("my_file.docx")
    editor.replace_text("old text", "new text")
    editor.add_comment("Important note", on_text="some phrase")
    editor.save("my_file_edited.docx")

FULL USAGE: See each method's docstring below.
"""

import os
import re
import copy
import zipfile
import shutil
import tempfile
import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
#  Namespace map used throughout
# ─────────────────────────────────────────────
NSMAP = {
    'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _w(tag):
    """Shorthand: returns '{namespace}tag' string."""
    return f'{{{W}}}{tag}'


# ══════════════════════════════════════════════════════════════════
#  DocxEditor  —  main class
# ══════════════════════════════════════════════════════════════════

class DocxEditor:
    """
    A complete .docx editor for AI projects.

    Supports:
    - Reading text content paragraph by paragraph
    - Simple find & replace
    - Tracked insertions and deletions (visible in Word's Review pane)
    - Inline comments
    - Bold / italic / underline / color formatting on any text
    - Adding paragraphs, headings, bullet lists
    - Adding and editing tables
    - Creating brand-new documents from scratch
    - Saving with a new filename (non-destructive)

    Parameters
    ----------
    filepath : str or None
        Path to an existing .docx file.  Pass None to create a new document.
    author : str
        Name shown in tracked changes and comments (default: "AI Assistant").
    """

    def __init__(self, filepath=None, author="AI Assistant"):
        self.author = author
        self._change_id = 1          # incrementing ID for tracked changes
        self._comment_id = 0         # incrementing ID for comments
        self._comments_xml = []      # accumulated <w:comment> elements

        if filepath is None:
            self.doc = Document()
            self.filepath = None
        else:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File not found: {filepath}")
            self.doc = Document(filepath)
            self.filepath = filepath

    # ──────────────────────────────────────────
    #  1. READ
    # ──────────────────────────────────────────

    def read_text(self) -> str:
        """Return all document text as a single string."""
        return "\n".join(p.text for p in self.doc.paragraphs)

    def read_paragraphs(self) -> list:
        """Return list of dicts: [{index, style, text}, ...]"""
        result = []
        for i, p in enumerate(self.doc.paragraphs):
            result.append({
                "index": i,
                "style": p.style.name,
                "text":  p.text,
            })
        return result

    def find_paragraph(self, search_text: str) -> list:
        """Return list of paragraph indices that contain search_text."""
        return [i for i, p in enumerate(self.doc.paragraphs)
                if search_text.lower() in p.text.lower()]

    def read_tables(self) -> list:
        """Return all table data as list of 2-D arrays (rows × cols)."""
        tables = []
        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return tables

    # ──────────────────────────────────────────
    #  2. SIMPLE REPLACE (no track changes)
    # ──────────────────────────────────────────

    def replace_text(self, old: str, new: str, case_sensitive=True) -> int:
        """
        Replace all occurrences of `old` with `new` in the document body.
        Works across runs inside each paragraph.
        Returns the number of replacements made.

        Example
        -------
        editor.replace_text("John", "Jane")
        """
        count = 0
        flags = 0 if case_sensitive else re.IGNORECASE
        for para in self.doc.paragraphs:
            full_text = para.text
            if not re.search(re.escape(old), full_text, flags):
                continue
            # Rebuild the paragraph with a single run to avoid split-run issues
            new_text = re.sub(re.escape(old), new, full_text, flags=flags)
            count += full_text.count(old) if case_sensitive else len(
                re.findall(re.escape(old), full_text, flags))
            # Clear existing runs and add one merged run
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
        return count

    # ──────────────────────────────────────────
    #  3. TRACKED CHANGES (insertions / deletions)
    # ──────────────────────────────────────────

    def insert_tracked(self, paragraph_index: int, insert_after: str, new_text: str):
        """
        Insert `new_text` as a tracked insertion right after `insert_after`
        inside the paragraph at `paragraph_index`.

        The insertion shows as green underlined text in Word's Review pane.

        Example
        -------
        editor.insert_tracked(0, "Hello", " World")
        # Paragraph 0: "Hello" → "Hello [inserted: World]"
        """
        para = self.doc.paragraphs[paragraph_index]
        p_elem = para._p
        date = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

        # Find the run containing insert_after
        for run in para.runs:
            if insert_after in run.text:
                r_elem = run._r
                # Split run at insertion point
                idx = run.text.index(insert_after) + len(insert_after)
                before = run.text[:idx]
                after  = run.text[idx:]
                run.text = before

                # Build <w:ins> element
                ins = OxmlElement('w:ins')
                ins.set(_w('id'),     str(self._change_id))
                ins.set(_w('author'), self.author)
                ins.set(_w('date'),   date)
                self._change_id += 1

                new_r = OxmlElement('w:r')
                if run._r.find(_w('rPr')) is not None:
                    new_r.append(copy.deepcopy(run._r.find(_w('rPr'))))
                t = OxmlElement('w:t')
                t.text = new_text
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_r.append(t)
                ins.append(new_r)
                r_elem.addnext(ins)

                # Remainder run (after insertion point)
                if after:
                    after_r = copy.deepcopy(r_elem)
                    after_r.find(_w('t')).text = after
                    ins.addnext(after_r)
                return True
        raise ValueError(f"Text '{insert_after}' not found in paragraph {paragraph_index}.")

    def delete_tracked(self, paragraph_index: int, delete_text: str):
        """
        Mark `delete_text` as a tracked deletion inside the paragraph.
        The text shows as red strikethrough in Word's Review pane.

        Example
        -------
        editor.delete_tracked(2, "unnecessary phrase")
        """
        para = self.doc.paragraphs[paragraph_index]
        date = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

        for run in para.runs:
            if delete_text in run.text:
                r_elem = run._r
                idx   = run.text.index(delete_text)
                before = run.text[:idx]
                after  = run.text[idx + len(delete_text):]
                run.text = before

                # Build <w:del>
                del_elem = OxmlElement('w:del')
                del_elem.set(_w('id'),     str(self._change_id))
                del_elem.set(_w('author'), self.author)
                del_elem.set(_w('date'),   date)
                self._change_id += 1

                del_r = OxmlElement('w:r')
                if r_elem.find(_w('rPr')) is not None:
                    del_r.append(copy.deepcopy(r_elem.find(_w('rPr'))))
                del_t = OxmlElement('w:delText')
                del_t.text = delete_text
                del_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                del_r.append(del_t)
                del_elem.append(del_r)
                r_elem.addnext(del_elem)

                if after:
                    after_r = copy.deepcopy(r_elem)
                    after_r.find(_w('t')).text = after
                    del_elem.addnext(after_r)
                return True
        raise ValueError(f"Text '{delete_text}' not found in paragraph {paragraph_index}.")

    def replace_tracked(self, paragraph_index: int, old_text: str, new_text: str):
        """
        Replace `old_text` with `new_text` as a tracked change:
        deletion of old + insertion of new, side by side.

        Example
        -------
        editor.replace_tracked(1, "30 days", "60 days")
        """
        self.delete_tracked(paragraph_index, old_text)
        # Re-find the deletion element and insert after it
        para = self.doc.paragraphs[paragraph_index]
        date = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        p_elem = para._p

        # Find the last <w:del> we just added and append <w:ins> after it
        del_elems = p_elem.findall(_w('del'))
        if not del_elems:
            raise RuntimeError("Could not locate the deletion element.")
        last_del = del_elems[-1]

        ins = OxmlElement('w:ins')
        ins.set(_w('id'),     str(self._change_id))
        ins.set(_w('author'), self.author)
        ins.set(_w('date'),   date)
        self._change_id += 1

        new_r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = new_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(t)
        ins.append(new_r)
        last_del.addnext(ins)

    # ──────────────────────────────────────────
    #  4. COMMENTS
    # ──────────────────────────────────────────

    def add_comment(self, comment_text: str, on_text: str, paragraph_index: int = None):
        """
        Add a comment bubble on `on_text` (first occurrence).
        Optionally restrict search to a specific paragraph_index.

        The comment appears in Word's Review pane, same as a human adding it.

        Example
        -------
        editor.add_comment("Consider simplifying this sentence.", on_text="hereby")
        """
        date  = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        cid   = self._comment_id
        self._comment_id += 1

        # Build the <w:comment> element to store in comments.xml later
        comment_elem = OxmlElement('w:comment')
        comment_elem.set(_w('id'),     str(cid))
        comment_elem.set(_w('author'), self.author)
        comment_elem.set(_w('date'),   date)
        comment_elem.set(_w('initials'), self.author[:2].upper())

        cp = OxmlElement('w:p')
        cr = OxmlElement('w:r')
        ct = OxmlElement('w:t')
        ct.text = comment_text
        cr.append(ct)
        cp.append(cr)
        comment_elem.append(cp)
        self._comments_xml.append(comment_elem)

        # Find the paragraph containing on_text
        search_range = ([self.doc.paragraphs[paragraph_index]]
                        if paragraph_index is not None
                        else self.doc.paragraphs)

        for para in search_range:
            if on_text not in para.text:
                continue
            p_elem = para._p
            for run in para.runs:
                if on_text not in run.text:
                    continue
                r_elem = run._r
                idx    = run.text.index(on_text)
                before = run.text[:idx]
                after  = run.text[idx + len(on_text):]
                run.text = before

                # <w:commentRangeStart>
                crs = OxmlElement('w:commentRangeStart')
                crs.set(_w('id'), str(cid))
                r_elem.addnext(crs)

                # The annotated run
                ann_r = OxmlElement('w:r')
                if r_elem.find(_w('rPr')) is not None:
                    ann_r.append(copy.deepcopy(r_elem.find(_w('rPr'))))
                ann_t = OxmlElement('w:t')
                ann_t.text = on_text
                ann_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                ann_r.append(ann_t)
                crs.addnext(ann_r)

                # <w:commentRangeEnd>
                cre = OxmlElement('w:commentRangeEnd')
                cre.set(_w('id'), str(cid))
                ann_r.addnext(cre)

                # <w:r><w:commentReference/></w:r>
                ref_r = OxmlElement('w:r')
                rpr   = OxmlElement('w:rPr')
                rs    = OxmlElement('w:rStyle')
                rs.set(_w('val'), 'CommentReference')
                rpr.append(rs)
                ref_r.append(rpr)
                ref   = OxmlElement('w:commentReference')
                ref.set(_w('id'), str(cid))
                ref_r.append(ref)
                cre.addnext(ref_r)

                # Remainder
                if after:
                    aft_r = copy.deepcopy(r_elem)
                    aft_r.find(_w('t')).text = after
                    ref_r.addnext(aft_r)
                return True
        raise ValueError(f"Text '{on_text}' not found in the document.")

    # ──────────────────────────────────────────
    #  5. FORMATTING
    # ──────────────────────────────────────────

    def format_text(self, search_text: str,
                    bold=False, italic=False, underline=False,
                    color_hex: str = None, font_size_pt: int = None):
        """
        Apply formatting to every run that contains `search_text`.

        Parameters
        ----------
        search_text  : text to find and format
        bold         : True/False
        italic       : True/False
        underline    : True/False
        color_hex    : e.g. "FF0000" for red
        font_size_pt : e.g. 14

        Example
        -------
        editor.format_text("Important:", bold=True, color_hex="C00000")
        """
        for para in self.doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    if bold      is not False: run.bold      = bold
                    if italic    is not False: run.italic    = italic
                    if underline is not False: run.underline = underline
                    if color_hex:
                        r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                        run.font.color.rgb = RGBColor(r, g, b)
                    if font_size_pt:
                        run.font.size = Pt(font_size_pt)

    def set_paragraph_alignment(self, paragraph_index: int, alignment: str):
        """
        Set alignment of a paragraph.
        alignment: "left" | "center" | "right" | "justify"
        """
        mapping = {
            "left":    WD_ALIGN_PARAGRAPH.LEFT,
            "center":  WD_ALIGN_PARAGRAPH.CENTER,
            "right":   WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment not in mapping:
            raise ValueError(f"alignment must be one of {list(mapping.keys())}")
        self.doc.paragraphs[paragraph_index].alignment = mapping[alignment]

    # ──────────────────────────────────────────
    #  6. ADD CONTENT
    # ──────────────────────────────────────────

    def add_paragraph(self, text: str, bold=False, italic=False,
                      font_size=11, color_hex=None):
        """Append a new paragraph to the document."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color_hex:
            r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(r, g, b)
        return p

    def add_heading(self, text: str, level: int = 1):
        """
        Add a heading paragraph.
        level: 1 (biggest) to 9 (smallest)

        Example
        -------
        editor.add_heading("Chapter 1: Introduction", level=1)
        """
        return self.doc.add_heading(text, level=level)

    def add_bullet(self, text: str):
        """Add a bullet-list item."""
        return self.doc.add_paragraph(text, style="List Bullet")

    def add_numbered_item(self, text: str):
        """Add a numbered list item."""
        return self.doc.add_paragraph(text, style="List Number")

    def add_page_break(self):
        """Insert a page break."""
        self.doc.add_page_break()

    def insert_paragraph_at(self, index: int, text: str, style="Normal"):
        """
        Insert a paragraph BEFORE the paragraph at `index`.

        Example
        -------
        editor.insert_paragraph_at(0, "CONFIDENTIAL", style="Normal")
        """
        from docx.oxml.ns import qn
        new_para = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_para.append(r)
        ref = self.doc.paragraphs[index]._p
        ref.addprevious(new_para)

    # ──────────────────────────────────────────
    #  7. TABLES
    # ──────────────────────────────────────────

    def add_table(self, headers: list, rows: list, style="Table Grid"):
        """
        Add a table with headers and data rows.

        Parameters
        ----------
        headers : list of column header strings, e.g. ["Name", "Score"]
        rows    : list of lists, e.g. [["Alice", "95"], ["Bob", "87"]]
        style   : Word table style name (default: "Table Grid")

        Example
        -------
        editor.add_table(
            headers=["Feature", "Status"],
            rows=[["Login", "Done"], ["Search", "In Progress"]]
        )
        """
        col_count = len(headers)
        table = self.doc.add_table(rows=1, cols=col_count)
        table.style = style

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True

        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
        return table

    def edit_table_cell(self, table_index: int, row: int, col: int, new_text: str):
        """
        Replace the text in a specific table cell.

        Example
        -------
        editor.edit_table_cell(0, row=1, col=2, new_text="Updated value")
        """
        cell = self.doc.tables[table_index].cell(row, col)
        cell.text = new_text

    def add_table_row(self, table_index: int, values: list):
        """
        Append a new row to an existing table.

        Example
        -------
        editor.add_table_row(0, ["Charlie", "92"])
        """
        table = self.doc.tables[table_index]
        row = table.add_row()
        for i, val in enumerate(values):
            row.cells[i].text = str(val)

    # ──────────────────────────────────────────
    #  8. SAVE
    # ──────────────────────────────────────────

    def save(self, output_path: str = None):
        """
        Save the document.

        If `output_path` is None, overwrites the original file.
        If comments were added, they are injected into comments.xml automatically.

        Example
        -------
        editor.save("document_edited.docx")
        """
        target = output_path or self.filepath
        if target is None:
            raise ValueError("No output path specified and no original file loaded.")

        # Save normally first
        self.doc.save(target)

        # Inject comments if any were added
        if self._comments_xml:
            self._write_comments(target)

        print(f"✅ Saved: {target}")
        return target

    def _write_comments(self, docx_path: str):
        """Internal: inject comments.xml into the .docx ZIP."""
        tmp = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, 'r') as zin, \
             zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:

            # Copy all existing files
            for item in zin.infolist():
                if item.filename != 'word/comments.xml':
                    zout.writestr(item, zin.read(item.filename))

            # Build comments.xml
            root = etree.Element(
                _w('comments'),
                nsmap={'w': W,
                       'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                       'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'}
            )
            for c in self._comments_xml:
                root.append(c)
            comments_bytes = etree.tostring(root, xml_declaration=True,
                                            encoding='UTF-8', standalone=True)
            zout.writestr('word/comments.xml', comments_bytes)

            # Patch [Content_Types].xml if needed
            ct_data = zin.read('[Content_Types].xml').decode('utf-8')
            if 'comments' not in ct_data:
                ct_data = ct_data.replace(
                    '</Types>',
                    '<Override PartName="/word/comments.xml" '
                    'ContentType="application/vnd.openxmlformats-officedocument'
                    '.wordprocessingml.comments+xml"/></Types>'
                )
                zout.writestr('[Content_Types].xml', ct_data.encode('utf-8'))

            # Patch word/_rels/document.xml.rels if needed
            rels_name = 'word/_rels/document.xml.rels'
            rels_data = zin.read(rels_name).decode('utf-8')
            if 'comments' not in rels_data:
                rels_data = rels_data.replace(
                    '</Relationships>',
                    '<Relationship Id="rIdComments" '
                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
                    'relationships/comments" Target="comments.xml"/>'
                    '</Relationships>'
                )
                zout.writestr(rels_name, rels_data.encode('utf-8'))

        os.replace(tmp, docx_path)


# ══════════════════════════════════════════════════════════════════
#  DocxCreator  —  create brand-new .docx from scratch
# ══════════════════════════════════════════════════════════════════

class DocxCreator:
    """
    Create a brand-new .docx document from Python.

    Example
    -------
    creator = DocxCreator(title="My Report", author="AI Assistant")
    creator.add_heading("Introduction", level=1)
    creator.add_paragraph("This report covers...")
    creator.add_table(
        headers=["Item", "Value"],
        rows=[["Accuracy", "92%"], ["Speed", "Fast"]]
    )
    creator.save("report.docx")
    """

    def __init__(self, title: str = "", author: str = "AI Assistant",
                 page_width_inches=8.5, page_height_inches=11,
                 margin_inches=1.0):
        self.doc = Document()
        self.author = author

        # Page size
        section = self.doc.sections[0]
        section.page_width   = int(page_width_inches  * 914400)
        section.page_height  = int(page_height_inches * 914400)
        section.top_margin    = int(margin_inches * 914400)
        section.bottom_margin = int(margin_inches * 914400)
        section.left_margin   = int(margin_inches * 914400)
        section.right_margin  = int(margin_inches * 914400)

        if title:
            p = self.doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading(self, text: str, level: int = 1,
                    bold: bool = True, color_hex: str = None):
        """Add a heading (level 1–6)."""
        p = self.doc.add_heading(text, level=level)
        if p.runs:
            run = p.runs[0]
            if bold is not None: run.bold = bold
            if color_hex:
                r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(r, g, b)
        return p

    def add_paragraph(self, text: str, bold=False, italic=False,
                      font_size=11, color_hex=None, align="left"):
        """Add a body paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color_hex:
            r, g, b = (int(color_hex[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(r, g, b)
        alignments = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                      "center": WD_ALIGN_PARAGRAPH.CENTER,
                      "right": WD_ALIGN_PARAGRAPH.RIGHT,
                      "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
        p.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        return p

    def add_bullet(self, text: str):
        """Add a bullet list item."""
        return self.doc.add_paragraph(text, style="List Bullet")

    def add_numbered_item(self, text: str):
        """Add a numbered list item."""
        return self.doc.add_paragraph(text, style="List Number")

    def add_table(self, headers: list, rows: list,
                  header_bg="2E4057", header_fg="FFFFFF"):
        """
        Add a styled table with colored headers.

        Example
        -------
        creator.add_table(
            headers=["Name", "Role", "Status"],
            rows=[["Alice", "Dev", "Active"], ["Bob", "PM", "Active"]]
        )
        """
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Style header cells
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            r, g, b = (int(header_fg[i2:i2+2], 16) for i2 in (0, 2, 4))
            run.font.color.rgb = RGBColor(r, g, b)
            # Background shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  header_bg)
            tcPr.append(shd)

        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
        return table

    def add_page_break(self):
        """Insert a page break."""
        self.doc.add_page_break()

    def add_horizontal_rule(self):
        """Add a thin horizontal line (paragraph border)."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def save(self, output_path: str):
        """Save the document to disk."""
        self.doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        return output_path


# ══════════════════════════════════════════════════════════════════
#  Convenience function: AI-friendly edit entry point
# ══════════════════════════════════════════════════════════════════

def ai_edit_document(filepath: str, instructions: list, output_path: str = None,
                     author: str = "AI Assistant") -> str:
    """
    Apply a list of editing instructions to a .docx file.
    Designed as a single entry point for an AI that receives edit tasks.

    Parameters
    ----------
    filepath     : path to input .docx
    instructions : list of dicts describing edits (see examples below)
    output_path  : where to save; defaults to adding "_edited" suffix
    author       : name shown in tracked changes

    Instruction formats
    -------------------
    {"action": "replace",         "old": "foo",   "new": "bar"}
    {"action": "replace_tracked", "paragraph": 0, "old": "foo", "new": "bar"}
    {"action": "insert_tracked",  "paragraph": 0, "after": "Hello", "text": " World"}
    {"action": "delete_tracked",  "paragraph": 0, "text": "foo"}
    {"action": "comment",         "text": "Check this", "on": "some phrase"}
    {"action": "format",          "find": "Title", "bold": True, "color": "1F3864"}
    {"action": "add_paragraph",   "text": "New content"}
    {"action": "add_heading",     "text": "New Section", "level": 2}
    {"action": "add_bullet",      "text": "New bullet"}

    Example
    -------
    result = ai_edit_document("report.docx", [
        {"action": "replace",        "old": "2023",     "new": "2024"},
        {"action": "replace_tracked","paragraph": 1,    "old": "draft", "new": "final"},
        {"action": "comment",        "text": "Update this figure", "on": "42%"},
        {"action": "add_heading",    "text": "Appendix", "level": 1},
    ])
    print("Saved to:", result)
    """
    editor = DocxEditor(filepath, author=author)
    if output_path is None:
        base, ext = os.path.splitext(filepath)
        output_path = base + "_edited" + ext

    for inst in instructions:
        action = inst.get("action", "").lower()
        try:
            if action == "replace":
                editor.replace_text(inst["old"], inst["new"])

            elif action == "replace_tracked":
                editor.replace_tracked(inst["paragraph"], inst["old"], inst["new"])

            elif action == "insert_tracked":
                editor.insert_tracked(inst["paragraph"], inst["after"], inst["text"])

            elif action == "delete_tracked":
                editor.delete_tracked(inst["paragraph"], inst["text"])

            elif action == "comment":
                para = inst.get("paragraph", None)
                editor.add_comment(inst["text"], on_text=inst["on"],
                                   paragraph_index=para)

            elif action == "format":
                editor.format_text(
                    inst["find"],
                    bold=inst.get("bold", False),
                    italic=inst.get("italic", False),
                    underline=inst.get("underline", False),
                    color_hex=inst.get("color"),
                    font_size_pt=inst.get("size"),
                )

            elif action == "add_paragraph":
                editor.add_paragraph(inst["text"],
                                     bold=inst.get("bold", False),
                                     italic=inst.get("italic", False),
                                     font_size=inst.get("size", 11),
                                     color_hex=inst.get("color"))

            elif action == "add_heading":
                editor.add_heading(inst["text"], level=inst.get("level", 1))

            elif action == "add_bullet":
                editor.add_bullet(inst["text"])

            elif action == "add_table":
                editor.add_table(inst["headers"], inst["rows"])

            else:
                print(f"⚠️  Unknown action: '{action}' — skipped.")

        except Exception as e:
            print(f"⚠️  Error on instruction {inst}: {e}")

    return editor.save(output_path)


# ══════════════════════════════════════════════════════════════════
#  DEMO  (run this file directly to see it work)
# ══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # ── Demo 1: Create a new document ──
    print("Creating demo document...")
    c = DocxCreator(title="AI Search Engine Report", author="AI Assistant")
    c.add_heading("Project Overview", level=1)
    c.add_paragraph("This document was created entirely by the AI editing module.")
    c.add_paragraph("It supports headings, paragraphs, bullets, and tables.")
    c.add_bullet("CLIP-based multimodal search")
    c.add_bullet("FAISS vector similarity")
    c.add_bullet("Streamlit UI")
    c.add_table(
        headers=["Feature",        "Status"],
        rows=[
            ["Text Search",        "Done"],
            ["Image Search",       "Done"],
            ["GPU Support",        "Pending"],
            ["Incremental Index",  "Planned"],
        ]
    )
    c.save("demo_created.docx")

    # ── Demo 2: Edit with tracked changes + comments ──
    print("\nEditing with tracked changes...")
    result = ai_edit_document("demo_created.docx", [
        {"action": "replace",        "old": "Pending",  "new": "In Progress"},
        {"action": "comment",        "text": "This table needs a Status column header fix.",
                                     "on": "Feature"},
        {"action": "add_heading",    "text": "Conclusion",  "level": 1},
        {"action": "add_paragraph",  "text": "All planned features will be prioritized."},
    ], output_path="demo_edited.docx")

    print(f"\nDemo complete. Files: demo_created.docx, demo_edited.docx")
