import os
import io
import re
import zipfile
import shutil
from docx import Document
from datetime import datetime
import sys

# Ensure UTF-8 for terminal
if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ── TURBO-SURGICAL ENGINE v16.0 Logic ──

def smart_surgical_edit(template_bytes, ai_output):
    try:
        tmp_dir = "temp_surgery_test"
        if os.path.exists(tmp_dir): shutil.rmtree(tmp_dir)
        os.makedirs(tmp_dir)
        
        with zipfile.ZipFile(io.BytesIO(template_bytes), 'r') as zip_ref:
            zip_ref.extractall(tmp_dir)

        author = "KALI AI ASSISTANT"
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        total_swaps = 0
        changes = re.findall(r'"([^"]+)"\s*->\s*"([^"]+)"', ai_output)
        
        target_files = [os.path.join(tmp_dir, 'word', 'document.xml')]

        for xml_path in target_files:
            if not os.path.exists(xml_path): continue
            with open(xml_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()

            for old, new in changes:
                old_xml = old.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                new_xml = new.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                
                def q_anchor(t):
                    words = t.split()
                    if not words: return re.escape(t)
                    return r'(?:<[^>]+>|\s)*'.join([re.escape(w) for w in words])

                pattern = q_anchor(old_xml)
                match = re.search(f'(<w:rPr>.*?</w:rPr>)?(?:<[^>]+>)*({pattern})', xml_content, re.DOTALL)
                
                if match:
                    style_xml = match.group(1) if match.group(1) else ""
                    full_match = match.group(0)
                    redline = (
                        f'<w:del w:id="{total_swaps*10}" w:author="{author}" w:date="{date_str}">'
                        f'<w:r>{style_xml}<w:delText>{old_xml}</w:delText></w:r></w:del>'
                        f'<w:ins w:id="{total_swaps*10+1}" w:author="{author}" w:date="{date_str}">'
                        f'<w:r>{style_xml}<w:t>{new_xml}</w:t></w:r></w:ins>'
                    )
                    xml_content = xml_content.replace(full_match, redline, 1)
                    total_swaps += 1

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(xml_content)

        bio = io.BytesIO()
        with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root, dirs, files in os.walk(tmp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    rel_path = os.path.relpath(full_path, tmp_dir)
                    zip_out.write(full_path, rel_path)
        
        shutil.rmtree(tmp_dir)
        return bio.getvalue() if total_swaps > 0 else None

    except Exception as e:
        print(f"FAILED: {str(e)}")
        return None

print("STEP 1: Creating Source Document...")
doc = Document()
doc.add_heading('KALI SYSTEM AUDIT REPORT', 0)
doc.add_paragraph('Current Status: Baseline Operational')
doc.add_paragraph('Assigned to: Project Manager Alpha')
doc.save('source_test.docx')

with open('source_test.docx', 'rb') as f:
    orig_bytes = f.read()

print("STEP 2: Designing AI Revision Plan...")
ai_plan = '"Baseline Operational" -> "Strategic Optima v2.0"\n"Project Manager Alpha" -> "Elite Executive Beta"'

print("STEP 3: Executing Structural Surgery...")
result_bytes = smart_surgical_edit(orig_bytes, ai_plan)

if result_bytes:
    with open('edited_result.docx', 'wb') as f:
        f.write(result_bytes)
    print("SUCCESS: Surgery Completed. 'edited_result.docx' generated.")
    doc_new = Document('edited_result.docx')
    all_text = " ".join([p.text for p in doc_new.paragraphs])
    print(f"VERIFICATION: Found text -> '{all_text}'")
    if "Strategic Optima" in all_text:
        print("RESULT: Data Verified. System Operational.")
else:
    print("ERROR: Surgery Failed to find targets.")
