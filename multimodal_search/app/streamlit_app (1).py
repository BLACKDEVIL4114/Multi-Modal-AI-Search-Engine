import streamlit as st
import os
from PIL import Image
import sys

# Fix pathing so search module and data are discoverable
root_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if root_path not in sys.path:
    sys.path.append(root_path)

from search.search_core import search_by_text, search_by_image
from utils.docx_handler import edit_docx_and_return

st.set_page_config(page_title="Multimodal Search + DOCX Editor", page_icon="🔍", layout="wide")


def extract_docx_text(docx_bytes) -> str:
    import tempfile
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name
    doc = Document(tmp_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text.strip())
    os.unlink(tmp_path)
    return "\n".join(lines[:80])


def ai_parse_instructions(user_instruction: str, doc_text: str, api_key: str) -> dict:
    import requests, json

    prompt = f"""You are a document editing assistant. The user has a DOCX file and wants to make changes to it.

Here is a sample of the document's current content:
---
{doc_text}
---

The user said: "{user_instruction}"

Your job: Figure out exactly what text needs to be replaced with what.
Return ONLY a valid JSON object like this (no explanation, no markdown, no extra text):
{{"replacements": {{"old text 1": "new text 1", "old text 2": "new text 2"}}}}

Rules:
- The "old text" must be the EXACT text currently in the document (copy it from the document content above)
- If the user says something vague like "fix my name", look at the document for a name and figure out what they mean
- If you cannot find what to replace, return {{"replacements": {{}}}}
- Only return the JSON. Nothing else."""

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 1000}
    }

    resp = requests.post(url, headers=headers, json=body)
    data = resp.json()

    if "error" in data:
        raise Exception(f"Gemini Error: {data['error'].get('message', str(data['error']))}")

    text = data["candidates"][0]["content"]["parts"][0]["text"]
    text = text.strip().replace("```json", "").replace("```", "").strip()
    parsed = json.loads(text)
    return parsed.get("replacements", {})


def render_docx_editor():
    st.title("📄 AI-Powered DOCX Editor")
    st.markdown("Upload your document and **just tell the AI what to change** — in plain English!")

    with st.expander("🔑 Google Gemini API Key (required for AI mode)", expanded=True):
        api_key = st.text_input("Enter your Gemini API Key", type="password", key="docx_api_key", placeholder="AIza...")
        st.caption("Free API key from aistudio.google.com — Your key is never stored.")

    uploaded_docx = st.file_uploader("📂 Upload a DOCX file", type=["docx"])

    doc_text = ""
    if uploaded_docx:
        doc_text = extract_docx_text(uploaded_docx.getvalue())
        with st.expander("👁️ Preview document content", expanded=False):
            st.text(doc_text)

    st.markdown("---")
    st.markdown("### ✏️ What do you want to change?")
    st.markdown("Just describe it naturally — like you're telling a friend:")

    instr_text = st.text_area(
        "Your editing instructions:",
        placeholder=(
            "Examples:\n"
            "• change my name from Henil Patel to Himanshu Kotval\n"
            "• replace the enrollment number 221130116048 with 22113006024\n"
            "• fix the date to 17 April 2026\n"
            "• or just: Henil Patel → Himanshu Kotval"
        ),
        height=150,
    )

    col1, col2 = st.columns(2)
    with col1:
        ai_btn = st.button("🤖 Let AI Understand & Edit", use_container_width=True, type="primary")
    with col2:
        manual_btn = st.button("⚡ Manual Mode (old → new)", use_container_width=True)

    if (ai_btn or manual_btn) and uploaded_docx is None:
        st.error("Please upload a DOCX file first.")
        return

    if not instr_text.strip():
        if ai_btn or manual_btn:
            st.warning("Please enter your editing instructions.")
        return

    replacements = {}

    if ai_btn:
        if not api_key:
            st.error("Please enter your Gemini API key above to use AI mode.")
            return
        with st.spinner("🤖 Gemini AI is reading your instructions..."):
            try:
                replacements = ai_parse_instructions(instr_text, doc_text, api_key)
            except Exception as e:
                st.error(f"AI parsing failed: {e}")
                return

    elif manual_btn:
        for line in instr_text.strip().splitlines():
            for sep in ["→", "->", "=>"]:
                if sep in line:
                    old, new = line.split(sep, 1)
                    replacements[old.strip()] = new.strip()
                    break

    if not replacements:
        st.warning("⚠️ No replacements detected. Try being more specific, or use manual mode with `old → new` format.")
        return

    st.markdown("### ✅ AI Detected These Replacements:")
    for old, new in replacements.items():
        st.markdown(f"- **`{old}`** → **`{new}`**")

    confirm = st.button("📝 Apply & Download Edited Document", use_container_width=True)

    if confirm:
        with st.spinner("Editing document..."):
            try:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
                    tmp_in.write(uploaded_docx.getvalue())
                    tmp_in_path = tmp_in.name

                tmp_out_path = tmp_in_path.replace(".docx", "_edited.docx")
                final_path = edit_docx_and_return(tmp_in_path, replacements, tmp_out_path)

                with open(final_path, "rb") as f:
                    st.download_button(
                        label="⬇️ DOWNLOAD EDITED DOCUMENT",
                        data=f,
                        file_name=f"EDITED_{uploaded_docx.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="docx_final_dl"
                    )
                st.success(f"✅ Done! Applied {len(replacements)} replacement(s).")
                st.balloons()
                if os.path.exists(tmp_in_path):
                    os.unlink(tmp_in_path)
            except Exception as exc:
                st.error(f"Error during editing: {exc}")


def main():
    workspace = st.sidebar.radio("Workspace", ["Image Search", "DOCX Editor"], index=0)

    if workspace == "DOCX Editor":
        render_docx_editor()
        return

    st.title("🔍 Multimodal Image Search Engine")
    st.markdown("Search for images using text descriptions or other images.")

    with st.sidebar:
        st.header("⚙️ Search Settings")
        top_k = st.slider("Number of results (Top-K)", min_value=1, max_value=20, value=5)
        st.divider()
        st.info("Performance: Using GPU acceleration if available.")

    index_path = os.path.join(root_path, "embeddings", "vector.index")
    if not os.path.exists(index_path):
        st.error(f"Please run `pipeline/build_index.py` first. Index not found at {index_path}")
        return

    search_mode = st.radio("Select Search Mode:", ("Search by Text", "Search by Image"), horizontal=True)
    results = []

    if search_mode == "Search by Text":
        query = st.text_input("Describe what you are looking for", placeholder="e.g. a cat on the floor")
        search_btn = st.button("Search")
        if (search_btn or query) and query:
            with st.spinner("Searching..."):
                results = search_by_text(query, top_k=top_k)
    else:
        uploaded_file = st.file_uploader("Upload an image", type=["jpg", "png", "jpeg"])
        if uploaded_file is not None:
            try:
                img = Image.open(uploaded_file)
                st.image(img, caption="Query Image", width=200)
                search_btn = st.button("Search")
                if search_btn:
                    with st.spinner("Searching..."):
                        results = search_by_image(img, top_k=top_k)
            except Exception as e:
                st.error(f"Could not open image: {e}")

    if results:
        st.subheader(f"Results (Showing top {len(results)})")
        cols = st.columns(3)
        for i, res in enumerate(results):
            with cols[i % 3]:
                img_display_path = res["image_path"]
                if not os.path.isabs(img_display_path):
                    img_display_path = os.path.join(root_path, img_display_path)
                if os.path.exists(img_display_path):
                    try:
                        display_img = Image.open(img_display_path)
                        st.image(display_img, use_column_width=True)
                        st.write(f"**Match: {int(res['score'] * 100)}%**")
                        st.caption(f"Path: {os.path.basename(img_display_path)}")
                    except:
                        st.error("Error loading image file.")
                else:
                    st.warning(f"File not found: {img_display_path}")
    elif "search_btn" in locals() and search_btn:
        st.info("No results found. Try a different query.")


if __name__ == "__main__":
    main()
