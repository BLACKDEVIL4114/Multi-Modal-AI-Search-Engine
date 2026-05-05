from docx import Document
import os
import shutil
from datetime import datetime

def load_doc(source):
    # Handle both paths (strings) and streams (BytesIO)
    if isinstance(source, str) and not os.path.exists(source):
        raise FileNotFoundError(f"File not found: {source}")
    return Document(source)

def verify_docx_bytes(data):
    """Verifies that a binary docx stream is architecturally sound."""
    import io
    try:
        source = io.BytesIO(data)
        Document(source)
        return True
    except Exception:
        return False

def save_doc(doc, path):
    # Backup before overwrite
    backup = None
    if os.path.exists(path):
        timestamp = datetime.now().strftime('%H%M%S')
        backup = path.replace(".docx", f"_backup_{timestamp}.docx")
        shutil.copy2(path, backup)
    
    try:
        doc.save(path)
        # Verify immediately
        Document(path)
    except Exception as e:
        # Restore backup if corrupt or failed
        if backup and os.path.exists(backup):
            shutil.copy2(backup, path)
            raise RuntimeError(f"Save produced corrupt file — backup restored: {str(e)}")
        else:
            raise RuntimeError(f"Critical Save Failure: {str(e)}")
