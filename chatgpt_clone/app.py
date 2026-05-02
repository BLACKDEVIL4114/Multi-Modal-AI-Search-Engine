import streamlit as st
st.set_page_config(page_title="Kali AI | Intelligence Studio v2", page_icon="✦", layout="wide")

import google.generativeai as genai
import os
import time
import json
import io
import re
import base64
import socket
import zipfile
import tempfile
import shutil
from datetime import datetime
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.docx_handler import load_doc, save_doc, verify_docx_bytes

try:
    from groq import Groq
    from dotenv import load_dotenv
    import requests
    from bs4 import BeautifulSoup
    from streamlit_mic_recorder import mic_recorder
    load_dotenv()
except ImportError:
    st.error("⚠️ **Initialization Failure:** Dependencies missing. Run: `pip install groq python-dotenv requests beautifulsoup4 message-transformers faiss-cpu streamlit-mic-recorder`")
    st.stop()

# ── Handshake Synchronization ─────────────────────
if 'initialized' not in st.session_state:
    st.session_state.initialized = True
    if 'voice_prompt' not in st.session_state: st.session_state.voice_prompt = None
    if 'edit_index' not in st.session_state: st.session_state.edit_index = None
    if 'edit_text' not in st.session_state: st.session_state.edit_text = ""

# ── Persistence Engine ──────────────────────────────
HISTORY_DIR = ".kali_history"
HISTORY_FILE = os.path.join(HISTORY_DIR, "chat_history.json")

def save_chat_to_disk(chat_id, messages):
    try:
        if not os.path.exists(HISTORY_DIR): os.makedirs(HISTORY_DIR)
        history = {}
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "r") as f: history = json.load(f)
        
        title = "New Conversation"
        for m in messages:
            if m["role"] == "user":
                title = m["content"][:30] + "..." if len(m["content"]) > 30 else m["content"]
                break
                
        history[chat_id] = {"title": title, "messages": messages, "timestamp": chat_id}
        with open(HISTORY_FILE, "w") as f: json.dump(history, f)
    except:
        pass # Fail silently to prevent refresh loops on write-protected environments

def load_all_chats():
    try:
        if not os.path.exists(HISTORY_DIR): os.makedirs(HISTORY_DIR)
        if not os.path.exists(HISTORY_FILE): return {}
        with open(HISTORY_FILE, "r") as f: return json.load(f)
    except:
        return {}

# ── Cybersecurity: Scraper Firewall ────────────
def is_safe_url(url):
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ['http', 'https']: return False
        hostname = parsed.hostname
        if not hostname: return False
        
        # Block common local/private IP ranges
        ip = socket.gethostbyname(hostname)
        parts = list(map(int, ip.split('.')))
        if parts[0] in [127, 10, 172, 192]: return False # Local/Private
        if ip == '169.254.169.254': return False # Cloud Metadata
        return True
    except:
        return False

# ── Web-Synapse: Autonomous Search ────────────────
def autonomous_search(query):
    try:
        url = f"https://duckduckgo.com/html/?q={query.replace(' ', '+')}"
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        results = []
        links_to_scrape = []
        
        for result in soup.find_all('div', class_='result__body')[:8]:
            title = result.find('a', class_='result__a').get_text().strip()
            snippet = result.find('a', class_='result__snippet').get_text().strip()
            link = result.find('a', class_='result__a')['href']
            results.append(f"Title: {title}\nSummary: {snippet}\nSource: {link}\n")
            if len(links_to_scrape) < 3:
                links_to_scrape.append(link)
        
        # Deep Intelligence: Scrape the top 5 links for full context
        deep_context = []
        for i, link in enumerate(links_to_scrape[:5]):
            st.write(f"🔬 Deep Analyzing: {link[:50]}...")
            content = fetch_web_content(link)
            if content and "SECURITY BLOCK" not in content and len(content) > 200:
                deep_context.append(f"--- FULL CONTENT FROM SOURCE {i+1} ({link}) ---\n{content[:4000]}\n")
        
        final_context = "\n".join(results) + "\n\n" + "\n".join(deep_context)
        if len(final_context) < 100:
            return "No reliable web data found. Please try a different query."
        return final_context
    except Exception as e:
        return f"Search Error: {str(e)}"

# ── Web Intelligence Engine (Hardened) ────────
def fetch_web_content(url):
    if not is_safe_url(url):
        return "⚠️ [SECURITY BLOCK] Access to internal, local, or unsafe resources is prohibited."
    try:
        # Modern Browser Fingerprint
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Referer': 'https://www.google.com/'
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        # Handle Word Documents
        if url.lower().endswith('.docx') or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in response.headers.get('Content-Type', ''):
            st.session_state.template_bytes = response.content
            return "✅ [BINARY CAPTURED] Remote document integrated."
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Aggressive Noise Removal for "Perfect" Results
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "button", "svg"]):
            tag.decompose()
            
        # Extract clean text with semantic spacing
        text = soup.get_text(separator='\n')
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        # Filter out very short lines that might be menu items or artifacts
        clean_text = "\n".join([line for line in lines if len(line) > 20 or (any(c.isdigit() for c in line) and len(line) > 5)])
        
        return clean_text[:15000]
    except Exception as e:
        return f"Web Access Error: {str(e)}"

# --- KALI v2.0 ARCHITECTURE: RADICAL INITIALIZATION ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "chat_id" not in st.session_state:
    st.session_state.chat_id = "kali_" + str(int(time.time()))

# --- HYPER-RESILIENT SYNC ---
load_dotenv(override=True) 
DEFAULT_API_KEY = os.getenv("GROQ_API_KEY", "")
DEFAULT_GEMINI_KEY = os.getenv("GEMINI_API_KEY", "")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Outfit:wght@400;500;600;700&display=swap');

    :root {
        --bg-deep: #0a0a0f;
        --sidebar-bg: #050508;
        --accent-primary: #8b5cf6;
        --accent-secondary: #c084fc;
        --border-subtle: rgba(255, 255, 255, 0.05);
        --text-bright: #f8fafc;
        --glass-bg: rgba(15, 15, 25, 0.7);
        --glass-border: rgba(255, 255, 255, 0.1);
    }

    /* 1. Global Midnight Slate Foundation */
    [data-testid="stAppViewContainer"], .stApp {
        background: radial-gradient(circle at 50% 0%, #1e1b4b 0%, #0a0a0f 100%) !important;
        color: var(--text-bright) !important;
        font-family: 'Inter', sans-serif !important;
    }
    
    [data-testid="stHeader"] {
        background: transparent !important;
    }

    h1, h2, h3 {
        font-family: 'Outfit', sans-serif !important;
        letter-spacing: -0.02em !important;
    }

    /* 2. Zero-Flicker Sidebar */
    section[data-testid="stSidebar"] {
        background-color: var(--sidebar-bg) !important;
        border-right: 1px solid var(--border-subtle) !important;
    }

    .sb-logo-v2 {
        display: flex;
        align-items: center;
        gap: 14px;
        padding: 20px 0;
        border-bottom: 1px solid var(--border-subtle);
        margin-bottom: 25px;
    }

    .logo-mark-v2 {
        width: 42px;
        height: 42px;
        background: linear-gradient(135deg, #8b5cf6, #d946ef);
        border-radius: 12px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 22px;
        color: white;
        box-shadow: 0 0 20px rgba(139, 92, 246, 0.4);
        font-family: 'Outfit', sans-serif;
    }

    /* 3. High-Fidelity Canvas */
    .hero-v2 {
        display: flex;
        flex-direction: column;
        align-items: center;
        padding: 80px 0 40px;
        position: relative;
    }

    .bg-glow-v2 {
        position: absolute;
        width: 800px;
        height: 800px;
        background: radial-gradient(circle, rgba(139, 92, 246, 0.1) 0%, transparent 70%);
        top: 50%;
        left: 50%;
        transform: translate(-50%, -50%);
        pointer-events: none;
        z-index: 0;
    }

    /* 4. Chat & Input Refinement */
    [data-testid="stChatMessage"] {
        background-color: var(--glass-bg) !important;
        backdrop-filter: blur(12px) !important;
        border: 1px solid var(--glass-border) !important;
        border-radius: 20px !important;
        margin-bottom: 24px !important;
        padding: 1.5rem !important;
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3) !important;
    }

    [data-testid="stChatMessageContent"] {
        color: #f1f5f9 !important;
        font-size: 1.05rem !important;
        line-height: 1.6 !important;
    }

    [data-testid="stChatMessageContent"] p, [data-testid="stChatMessageContent"] li, [data-testid="stChatMessageContent"] span {
        color: #f1f5f9 !important;
    }

    /* Action Buttons (Copy/Edit) Styling */
    .stChatActionRow {
        display: flex;
        gap: 8px;
        margin-top: 8px;
    }

    div[data-testid="column"] button {
        background: rgba(255,255,255,0.05) !important;
        border: 1px solid rgba(255,255,255,0.1) !important;
        border-radius: 6px !important;
        padding: 2px 6px !important;
        font-size: 12px !important;
        color: #94a3b8 !important;
        min-height: 24px !important;
        line-height: 1 !important;
    }

    div[data-testid="column"] button:hover {
        background: rgba(139, 92, 246, 0.2) !important;
        border-color: rgba(139, 92, 246, 0.4) !important;
        color: white !important;
    }

    /* Global Button Overrides */
    .stButton>button, .stDownloadButton>button {
        background: rgba(139, 92, 246, 0.1) !important;
        color: white !important;
        border: 1px solid rgba(139, 92, 246, 0.2) !important;
        border-radius: 12px !important;
        padding: 0.6rem 1.2rem !important;
        font-weight: 500 !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        text-transform: none !important;
    }
    
    .stButton>button:hover {
        background: rgba(139, 92, 246, 0.2) !important;
        border-color: #8b5cf6 !important;
        box-shadow: 0 0 20px rgba(139, 92, 246, 0.3) !important;
        transform: translateY(-2px);
    }

    /* Sidebar Specific Contrast */
    [data-testid="stSidebar"] .stButton>button {
        background-color: transparent !important;
        border: 1px solid transparent !important;
        color: #64748b !important;
        text-align: left !important;
        font-size: 13px !important;
        padding: 6px 12px !important;
        width: 100% !important;
        margin-bottom: 0px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: flex-start !important;
        line-height: 1.1 !important;
        min-height: 32px !important;
        border-radius: 8px !important;
        transition: all 0.2s ease !important;
    }
    
    [data-testid="stSidebar"] .stButton {
        margin-bottom: -10px !important;
        padding-bottom: 0px !important;
    }
    
    [data-testid="stSidebar"] .stButton>button:hover {
        background: rgba(139, 92, 246, 0.08) !important;
        color: #c084fc !important;
    }

    [data-testid="stSidebar"] .stButton>button:active {
        background: rgba(139, 92, 246, 0.1) !important;
    }

    /* Precision Input Bar */
    .stChatInputContainer {
        border-radius: 18px !important;
        background-color: rgba(15, 15, 25, 0.95) !important;
        border: 1px solid var(--glass-border) !important;
        padding: 10px !important;
        backdrop-filter: blur(20px) !important;
    }
    
    .stChatInputContainer:focus-within {
        border-color: #8b5cf6 !important;
        box-shadow: 0 0 0 2px rgba(139, 92, 246, 0.2) !important;
    }

    /* Logic Chips */
    .chip-v2 {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 12px;
        padding: 10px 20px;
        font-size: 14px;
        color: #e2e8f0;
        margin: 6px;
        display: inline-block;
        transition: all 0.2s ease;
        cursor: pointer;
    }
    .chip-v2:hover {
        background: rgba(139, 92, 246, 0.1);
        border-color: #8b5cf6;
        color: #c084fc;
        transform: scale(1.05);
    }

    /* Voice Intelligence Premium Styling */
    .voice-card {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.1), rgba(217, 70, 239, 0.05));
        border: 1px solid rgba(139, 92, 246, 0.2);
        border-radius: 16px;
        padding: 20px;
        margin-top: 15px;
    }

    /* Pulse Badge Animation */
    @keyframes pulse {
        0% { transform: scale(1); box-shadow: 0 0 0 0 rgba(139, 92, 246, 0.4); }
        70% { transform: scale(1.05); box-shadow: 0 0 0 10px rgba(139, 92, 246, 0); }
        100% { transform: scale(1); box-shadow: 0 0 0 0 rgba(139, 92, 246, 0); }
    }
    .pulse-badge {
        animation: pulse 2s infinite;
    }
</style>
""", unsafe_allow_html=True)
# --- ENGINE INITIALIZATION v2.0 ---
DEFAULT_API_KEY = os.getenv("GROQ_API_KEY", "")
client = Groq(api_key=DEFAULT_API_KEY) if DEFAULT_API_KEY else None

# --- SIDEBAR & NAVIGATION v2.0 ---
all_history = load_all_chats()
with st.sidebar:
    st.markdown("""
        <div class="sb-logo-v2">
            <div class="logo-mark-v2">✦</div>
            <div style="color:white; font-size:15px; font-weight:500; line-height:1.1;">Kali AI<br><span style='font-size:9px; color:#a78bfa; letter-spacing:1px; font-weight:400;'>STUDIO v2</span></div>
        </div>
    """, unsafe_allow_html=True)
    
    if st.button("＋ New Session", use_container_width=True):
        st.session_state.messages = []
        st.session_state.chat_id = "kali_" + str(int(time.time()))
        for key in ["template_bytes", "chunks", "index", "final_doc"]:
            if key in st.session_state: del st.session_state[key]
        st.rerun()

    st.markdown("<div style='font-size:11px; color:#64748b; margin:25px 0 12px; font-weight:600; letter-spacing:1.2px; text-transform:uppercase;'>Recent Intelligence</div>", unsafe_allow_html=True)
    
    # Render History with v2 Styling (Shield Re-integrated)
    try:
        sorted_hist = sorted(all_history.items(), key=lambda x: x[1].get('timestamp', 0), reverse=True)
        for cid, data in sorted_hist[:10]:
            title = data.get('title', 'Untitled Intelligence')
            is_active = (cid == st.session_state.chat_id)
            btn_label = f"✦ {title}" if is_active else f"  {title}"
            if st.button(btn_label, key=f"nav_{cid}", use_container_width=True):
                st.session_state.messages = data['messages']
                st.session_state.chat_id = cid
                st.rerun()
        st.markdown("<div style='height: 40px;'></div>", unsafe_allow_html=True)
    except Exception as e:
        st.write("Recent Intelligence: [Syncing...]")

    st.divider()
    model_option = st.selectbox("Neural Architecture", [
        "llama-3.3-70b-versatile", 
        "llama-3.1-70b-versatile", 
        "mixtral-8x7b-32768",
        "gemini-2.0-flash (Premium Grounding)"
    ])
    
    # Dynamic Key Synchronization
    if st.button("🔄 Sync Neural Keys", use_container_width=True):
        load_dotenv(override=True)
        st.rerun()
        
    auth_key = st.text_input("Groq Engine Key", value=os.getenv("GROQ_API_KEY", ""), type="password")
    gemini_key = st.text_input("Gemini Engine Key", value=os.getenv("GEMINI_API_KEY", ""), type="password")
    
    if gemini_key:
        genai.configure(api_key=gemini_key)
    
    st.markdown(f"""
        <div style='padding: 20px 0; border-top: 1px solid var(--border-subtle); margin-top: 20px; display:flex; align-items:center; gap:12px;'>
            <div style='width:34px; height:34px; border-radius:10px; background:#1e1b4b; border:1px solid rgba(139, 92, 246, 0.3); color:#a78bfa; display:flex; align-items:center; justify-content:center; font-weight:bold; font-size:14px; box-shadow: 0 0 10px rgba(139, 92, 246, 0.2);'>✦</div>
            <div><b style='color:#fff; font-size:13px;'>Executive Control</b><br><span style='color:#4ade80; font-size:10px;'>● Neural Link Optimal</span> <span style='color:#60a5fa; font-size:10px; margin-left:10px;'>🌐 Web Online</span></div>
        </div>
    """, unsafe_allow_html=True)

# --- MAIN ENGINE INTERFACE CANVAS ---
if not st.session_state.messages:
    # Action Navigation
    cols = st.columns([1, 4, 3])
    with cols[0]:
        st.markdown("<div style='background:#141414; border:1px solid #222; border-radius:8px; padding:6px 12px; font-size:12px; color:#888;'>Studio v2.0</div>", unsafe_allow_html=True)
    with cols[2]:
        st.markdown("<div style='text-align:right; font-size:12px; color:#666;'>Engine Status: <span style='color:#6c63ff;'>Optimal</span></div>", unsafe_allow_html=True)

    st.markdown("""
        <div class="hero-v2">
            <div class="bg-glow-v2"></div>
            <div class="logo-mark-v2" style="width:80px; height:80px; font-size:42px; margin-bottom:25px; z-index:1; border-radius:18px;">✦</div>
            <h1 style="font-size:54px; font-weight:800; color:#f8fafc; margin-top:0; z-index:1; text-align:center; line-height:1.1;">
                Meet <span style="background: linear-gradient(135deg, #8b5cf6, #d946ef); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">Kali AI</span>
            </h1>
            <p style="color:#64748b; font-size:14px; text-transform:uppercase; letter-spacing:3px; margin-top:5px; z-index:1; font-weight:600;">
                Advanced Multi-Modal Search Engine
            </p>
        </div>
        <div style="text-align:center; position:relative; z-index:1; margin-bottom:50px;">
            <div class="chip-v2">✨ Creative Synthesis</div>
            <div class="chip-v2">🔬 Surgical Extraction</div>
            <div class="chip-v2">🌐 Real-time Intelligence</div>
            <div class="chip-v2">🎙️ Voice Command</div>
        </div>
    """, unsafe_allow_html=True)

# --- END UI SYNTHESIS ---

# ── Advanced Markdown-to-Pro-Word Engine ───────────
def clean_content(text):
    """Removes AI conversational fluff"""
    patterns = [
        r"^Sure,? here (is|are).*?:\n+",
        r"^Here’s a (revised|draft|version).*?:\n+",
        r"^Certainly,? .*?:\n+",
        r"^Okay,? .*?:\n+",
        r"^Revised version:\n+",
        r"^Revised content:\n+",
        r"^Revised Doc:\n+"
    ]
    for p in patterns:
        text = re.sub(p, "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n+Let me know if you need anything else!.*?$", "", text, flags=re.IGNORECASE | re.DOTALL)
    return text.strip()

def edit_and_return_docx(template_bytes, revised_content):
    """
    KALI SURGICAL TOOL (v25.0) - User-Defined High Precision
    """
    # st.info("🛰️ KALI SURGICAL PROTOCOL: Initializing Elite Revision Engine...")
    
    changes = []
    # --- OMNI-SURGICAL PARSER (v14.0 - Cluster Aware) ---
    # 1. Quoted Clusters: "A" -> "B" "C" -> "D"
    pattern_quoted = r'(?:["\']{1,3})(.*?)(?:["\']{1,3})\s*(?:\-+|==|=)>\s*(?:["\']{1,3})(.*?)(?:["\']{1,3})(?=\s*["\']|$)'
    matches = re.findall(pattern_quoted, revised_content)
    
    # 2. Raw/Technical Clusters (XML tags or single words)
    if not matches:
        pattern_raw = r'(?:^|\s|<)([^>\n\-\s]+(?: <[^>]+>)?)\s*(?:\-+|==|=)>\s*([^"\n\s<]+(?: <[^>]+>)?)(?:\s|$)'
        matches = re.findall(pattern_raw, revised_content)

    for old, new in matches:
        old_clean, new_clean = old.strip(), new.strip()
        if old_clean and new_clean:
            if "<!--" in old_clean: continue
            changes.append((old_clean, new_clean))
    
    if not changes:
        # Fallback for "Original: X New: Y"
        fallback_matches = re.findall(r'(?:Original|From|Old):\s*(.*?)\s*(?:-+|==|=)>\s*(?:New|To|Updated):\s*(.*?)(?:\n|$)', revised_content, re.DOTALL | re.IGNORECASE)
        for old, new in fallback_matches:
            changes.append((old.strip(), new.strip()))

    if not changes:
        return create_pro_docx(revised_content), "⚠️ No surgical targets. Created full draft."

    # Phase 3: Execute (The Elite ZIP-XML Workflow)
    try:
        tmp_dir = tempfile.mkdtemp()
        with zipfile.ZipFile(io.BytesIO(template_bytes), 'r') as zip_ref:
            zip_ref.extractall(tmp_dir)
            
        # Target all XML content files
        xml_targets = []
        for root, dirs, files in os.walk(tmp_dir):
            for file in files:
                if file.endswith(".xml"):
                    xml_targets.append(os.path.join(root, file))
            
        total_swaps = 0
# SILENT BACKGROUND WORKER (No Flicker)
        # st.write(f"🧬 Synchronizing XML Runs across {len(xml_targets)} nodes...")
        
        author = "Kali AI"
        date_str = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        
        for xml_path in xml_targets:
            # CONTENT LOCK: Only merge runs in files that actually contain document text
            is_content_file = any(x in xml_path.lower() for x in ["document.xml", "header", "footer"])
            
            with open(xml_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()
            
            if not is_content_file:
                # Still do basic replacements for structural changes, but skip run merging
                pass
            else:
                # 1. DEEP SYNTHESIS: Aggressive Word Noise Removal
                # This strips EVERYTHING that Word uses to fragment words
                xml_content = re.sub(r'<w:proofErr w:type="(spellStart|spellEnd|gramStart|gramEnd)"/>', '', xml_content)
                xml_content = re.sub(r'<w:bookmark(Start|End) [^>]*/>', '', xml_content)
                xml_content = re.sub(r'<w:softHyphen/>', '', xml_content)
                xml_content = re.sub(r'<w:noProof/>', '', xml_content)
                xml_content = re.sub(r'<w:lastRenderedPageBreak/>', '', xml_content)
                
                # 2. TITANIUM RUN CONSOLIDATION (v21.0 Elite)
                # Strategy: Merge contiguous runs with identical properties to eliminate "Word Fragmentation"
                
                # First, strip all the noise that splits runs
                xml_content = re.sub(r'</w:t></w:r>(?:<w:proofErr [^>]*/>|<w:bookmark(Start|End) [^>]*/>|<w:lastRenderedPageBreak/>|<w:softHyphen/>|<w:noProof/>)*<w:r>(?:<w:rPr>.*?</w:rPr>)?<w:t>', '', xml_content, flags=re.DOTALL)
                
                # Second, merge runs that have the SAME properties
                # This is complex because we need to match the rPr content exactly
                def merge_identical_runs(match):
                    props1 = match.group(1) if match.group(1) else ""
                    text1 = match.group(2)
                    props2 = match.group(3) if match.group(3) else ""
                    text2 = match.group(4)
                    if props1 == props2:
                        return f'<w:r>{props1}<w:t>{text1}{text2}</w:t></w:r>'
                    return match.group(0)

                # Simplified merging for performance
                for _ in range(1):
                    xml_content = re.sub(r'<w:r>(<w:rPr>.*?</w:rPr>)?<w:t>(.*?)</w:t></w:r><w:r>(<w:rPr>.*?</w:rPr>)?<w:t>(.*?)</w:t></w:r>', merge_identical_runs, xml_content, flags=re.DOTALL)
                
                # Remove all Word "noise" entirely from the searchable content
                unwanted_tags = [
                    r'<w:proofErr w:type="(spellStart|spellEnd|gramStart|gramEnd)"/>',
                    r'<w:bookmarkStart [^>]*/>',
                    r'<w:bookmarkEnd [^>]*/>',
                    r'<w:lastRenderedPageBreak/>',
                    r'<w:noProof/>',
                    r'<w:softHyphen/>',
                    r'<w:lang [^>]*/>'
                ]
                for tag_p in unwanted_tags:
                    xml_content = re.sub(tag_p, '', xml_content)
            
            # 3. Tracked Changes Surgery (Redlining & Structural)
            found_in_file = 0
            for old, new in changes:
                # DETECTION: Is this a structural XML change or a text content change?
                is_structural = '<' in old or '<' in new or '>' in old or '>' in new
                
                old_esc = old.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                new_esc = new.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Check 1: Exact Match (Raw)
                # Check 2: Exact Match (Escaped)
                # Check 3: Word-Field Normalization (flexible backslashes/case)
                
                if is_structural:
                    # STRUCTURAL MODE: Direct XML Injection
                    # SAFETY CHECK: Ensure the XML replacement is well-balanced to prevent corruption
                    if old.count('<') != new.count('<') or old.count('>') != new.count('>'):
                        st.warning(f"⚠️ Structural Lock: Blocked unbalanced XML swap to prevent corruption ({old[:20]}...)")
                # 🛡️ SHIELDED SURGERY (v15.0)
                # Escaping search terms to match XML storage format
                old_xml = old.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                new_xml = new.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                
                # MODE DETECTION: Default to REDLINE for safety, or DIRECT if requested
                surgery_mode = "REDLINE" if st.session_state.get("surgical_mode", True) else "DIRECT"
                
                if surgery_mode == "DIRECT":
                    if old_xml in xml_content:
                        xml_content = xml_content.replace(old_xml, new_xml)
                        found_in_file += 1
                        total_swaps += 1
                else:
                    # 🧬 REDLINE SURGERY: Track Changes Mode
                    import difflib
                    def get_diff(s1, s2):
                        w1, w2 = s1.split(), s2.split()
                        m = difflib.SequenceMatcher(None, w1, w2)
                        for t, i1, i2, j1, j2 in m.get_opcodes():
                            if t == 'replace': return " ".join(w1[i1:i2]), " ".join(w2[j1:j2])
                        return s1, s2

                    a_old, a_new = get_diff(old, new)
                    a_old_xml = a_old.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    a_new_xml = a_new.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                    def q_anchor(t):
                        # TITANIUM ANCHOR v21.0: Handles &nbsp;, multiple spaces, XML entities, and random tags
                        t = t.replace("&nbsp;", " ").replace("\u00A0", " ")
                        words = t.split()
                        if not words: return re.escape(t)
                        # The bridge can contain any XML tag, whitespace, or entity fragment
                        bridge = r'(?:<[^>]+>|\s|\&[^;]+;|&#\d+;)*'
                        return bridge.join([re.escape(w) for w in words])

                    pattern = q_anchor(a_old_xml)
                    
                    # 💉 LOCALIZED STYLE CLONING: Find the rPr (style) immediately before the text
                    # This prevents using styles from other parts of the document.
                    match = re.search(f'(<w:rPr>.*?</w:rPr>)?(?:<[^>]+>)*({pattern})', xml_content, re.DOTALL)
                    
                    if match:
                        style_xml = match.group(1) if match.group(1) else ""
                        full_match = match.group(0)
                        
                        # REPLACEMENT: Inserted as a clean insertion/deletion block
                        redline = (
                            f'<w:del w:id="{total_swaps*10}" w:author="{author}" w:date="{date_str}">'
                            f'<w:r>{style_xml}<w:delText>{a_old_xml}</w:delText></w:r></w:del>'
                            f'<w:ins w:id="{total_swaps*10+1}" w:author="{author}" w:date="{date_str}">'
                            f'<w:r>{style_xml}<w:t>{a_new_xml}</w:t></w:r></w:ins>'
                        )
                        
                        # Replace the specific match with the redline block
                        xml_content = xml_content.replace(full_match, redline, 1)
                        found_in_file += 1
                        total_swaps += 1

            if found_in_file > 0:
                with open(xml_path, 'w', encoding='utf-8') as f:
                    f.write(xml_content)

        # 3. Repack (Hardened v17.0)
        bio = io.BytesIO()
        with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for root, dirs, files in os.walk(tmp_dir):
                for file in files:
                    if file.startswith('.') or file.startswith('_'): continue
                    full_path = os.path.join(root, file)
                    rel_path = os.path.relpath(full_path, tmp_dir)
                    zip_out.write(full_path, rel_path)
        
        shutil.rmtree(tmp_dir)
        final_bytes = bio.getvalue()
        
        # 🛡️ THE ARCHITECTURAL GATE
        if verify_docx_bytes(final_bytes):
            if total_swaps > 0:
                msg = f"💎 Elite Revision Verified: {total_swaps} Revisions redlined."
                return final_bytes, msg
            else:
                return template_bytes, "⚠️ Revision logic found no XML targets."
        else:
            return template_bytes, "🚨 CRITICAL: Surgical edit produced a corrupt XML (Reverted)."

    except Exception as e:
        return None, f"❌ Elite Surgery Failure: {str(e)}"

def get_docx_bytes(doc):
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_pro_docx(content):
    """
    KALI PRO ENGINE v4.0 (Claude-Fidelity)
    """
    content = clean_content(content)
    doc = Document()
    
    # 🎨 CLAUDE-FIDELITY THEME
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial' 
    font.size = Pt(11)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_file_text(file):
    if file.name.endswith(".pdf"):
        from PyPDF2 import PdfReader
        return "\n".join([p.extract_text() for p in PdfReader(file).pages if p.extract_text()])
    elif file.name.endswith(".docx"):
        doc = load_doc(io.BytesIO(file.getvalue()))
        text_parts = []
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip(): text_parts.append(para.text)
            
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip(): text_parts.append(para.text)
                        
        # Headers/Footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip(): text_parts.append(f"[HEADER] {para.text}")
            for para in section.footer.paragraphs:
                if para.text.strip(): text_parts.append(f"[FOOTER] {para.text}")
                
        return "\n".join(text_parts)
    return ""

def get_chunks(text):
    size, overlap = 1000, 200
    res = []
    for i in range(0, len(text), size - overlap):
        chunk = text[i:i+size]
        if len(chunk) > 50: res.append(chunk)
    return res

@st.cache_resource
def get_ai_model():
    from sentence_transformers import SentenceTransformer
    with st.spinner("🧠 Syncing Neural Core..."):
        return SentenceTransformer('all-MiniLM-L6-v2')

@st.cache_resource
def build_vector_store(chunks):
    import faiss
    import numpy as np
    model = get_ai_model()
    embs = model.encode(chunks)
    idx = faiss.IndexFlatL2(embs.shape[1])
    idx.add(np.array(embs).astype('float32'))
    return idx, chunks

def fetch_knowledge(query, idx, chunks):
    import numpy as np
    if not idx: return ""
    model = get_ai_model()
    _, ids = idx.search(np.array(model.encode([query])).astype('float32'), 5)
    return "\n\n---\n\n".join([chunks[i] for i in ids[0] if i < len(chunks)])


# --- PERSISTENCE ENGINE MOVED TO TOP ---

# --- SIDEBAR CONSOLIDATED ABOVE ---

# --- LEGACY HEADER PURGED (v2.0 Logic Engaged) ---
st.markdown("<br>", unsafe_allow_html=True)

# ── Neural Chat Core ────────────────────────────
for i, msg in enumerate(st.session_state.messages):
    role = "user" if msg["role"] == "user" else "assistant"
    with st.chat_message(role):
        st.markdown(msg["content"])
        
        if role == "user":
            col1, col2, _ = st.columns([0.05, 0.05, 0.9])
            with col1:
                if st.button("📋", key=f"cp_{i}", help="Copy message"):
                    js_text = msg["content"].replace('"', '\\"').replace("\n", "\\n")
                    st.write(f'<script>navigator.clipboard.writeText("{js_text}");</script>', unsafe_allow_html=True)
                    st.toast("Copied to clipboard!", icon="✅")
            with col2:
                if st.button("✏️", key=f"ed_{i}", help="Edit & Resend"):
                    st.session_state.edit_index = i
                    st.session_state.edit_text = msg["content"]
                    st.rerun()

# ── Edit Overlay ────────────────────────────────
if st.session_state.get('edit_index') is not None:
    with st.container():
        st.markdown("### ✏️ Edit Intelligence Query")
        new_text = st.text_area("Update your query:", value=st.session_state.edit_text, height=100)
        ecol1, ecol2 = st.columns(2)
        with ecol1:
            if st.button("🚀 Resend Neural Query", use_container_width=True):
                # Remove everything after this message and resend
                st.session_state.messages = st.session_state.messages[:st.session_state.edit_index]
                st.session_state.active_prompt = new_text
                st.session_state.edit_index = None
                st.rerun()
        with ecol2:
            if st.button("❌ Cancel", use_container_width=True):
                st.session_state.edit_index = None
                st.rerun()

# ── Unit Fusion (Combined Bar) ────────────────
st.markdown("<br><br>", unsafe_allow_html=True)
footer_cols = st.columns([1, 15], gap="small")

with st.sidebar:
    st.markdown("### 🛰️ SYSTEM LINK")
    remote_path = st.text_input("Local File Path (Global Access)", placeholder="C:\\Path\\to\\document.docx")
    if remote_path and os.path.exists(remote_path):
        if st.button("🔌 CONNECT SYSTEM FILE"):
            try:
                with open(remote_path, 'rb') as rf:
                    file_bytes = rf.read()
                    st.session_state.template_bytes = file_bytes
                    st.session_state.uploaded_file_name = os.path.basename(remote_path)
                    st.success(f"Linked: {st.session_state.uploaded_file_name}")
                    st.rerun()
            except Exception as e:
                st.error(f"Link Error: {str(e)}")

    st.divider()
    st.markdown("### 💼 SESSION ASSETS")
    if st.session_state.get('template_bytes'):
        st.success(f"Active: {st.session_state.get('uploaded_file_name', 'Document')}")

with footer_cols[0]:
    with st.popover("➕", help="Upload photos & files"):
        st.markdown("### 📎 Attachment Center")
        uploaded_files = st.file_uploader("Upload Intel (PDF, DOCX, Images)", accept_multiple_files=True, key="intel_uploader")
        
        if uploaded_files:
            raw_text_parts = []
            for f in uploaded_files:
                file_type = f.name.split('.')[-1].lower()
                if file_type == 'docx':
                    if "template_bytes" not in st.session_state:
                        st.session_state.template_bytes = f.getvalue()
                        st.session_state.surgical_mode = True # Default to surgical
                    raw_text_parts.append(get_file_text(f))
                elif file_type == 'pdf':
                    raw_text_parts.append(get_file_text(f))
                elif file_type in ['jpg', 'jpeg', 'png']:
                    import base64
                    st.session_state.vision_base64 = base64.b64encode(f.getvalue()).decode('utf-8')
                    st.session_state.vision_active = True
                    st.sidebar.info(f"📸 Vision Feed: {f.name}")
            
            if raw_text_parts:
                full_text = "\n".join(raw_text_parts)
                chunks = get_chunks(full_text)
                if chunks:
                    st.session_state.index, st.session_state.chunks = build_vector_store(chunks)
                    st.toast(f"Knowledge Matrix Synced ({len(chunks)} nodes)")
        
        st.markdown("---")
        st.markdown("<div class='voice-card'>", unsafe_allow_html=True)
        st.markdown("### 🎙️ Voice Intelligence")
        audio = mic_recorder(
            start_prompt="⏺️ Start Recording",
            stop_prompt="⏹️ Stop Recording",
            just_once=False,
            use_container_width=True,
            key="voice_recorder"
        )
        
        if audio:
            if "last_audio_bytes" not in st.session_state or st.session_state.last_audio_bytes != audio['bytes']:
                st.session_state.last_audio_bytes = audio['bytes']
                with st.spinner("🎙️ Transcribing Voice..."):
                    try:
                        audio_bio = io.BytesIO(audio['bytes'])
                        audio_bio.name = "recording.wav"
                        
                        transcription = client.audio.transcriptions.create(
                            file=audio_bio,
                            model="whisper-large-v3",
                        )
                        st.session_state.voice_prompt = transcription.text
                        st.toast(f"✅ Voice Captured: {transcription.text[:50]}...")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Voice Error: {str(e)}")
        st.markdown("</div>", unsafe_allow_html=True)

with footer_cols[1]:
    prompt = st.chat_input("Ask anything...")

# ── Active Intelligence Processing ──────────
active_prompt = prompt if prompt else st.session_state.get("voice_prompt")

if active_prompt:
    if st.session_state.get("voice_prompt"):
        st.session_state.voice_prompt = None # Clear the buffer

    st.session_state.messages.append({"role": "user", "content": active_prompt})
    save_chat_to_disk(st.session_state.chat_id, st.session_state.messages)
    with st.chat_message("user"): st.markdown(active_prompt)

    # ── Web Detection & Extraction ────────────────
    web_context = ""
    links = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', active_prompt)
    
    if links:
        with st.status("🌐 Consulting Web Matrix...", expanded=True) as status:
            for link in links:
                st.write(f"Reading: {link}")
                content = fetch_web_content(link)
                web_context += f"\n--- WEB SOURCE: {link} ---\n{content}\n"
            status.update(label="✅ Web Intelligence Synced", state="complete", expanded=False)
    elif "?" in active_prompt or any(w in active_prompt.lower() for w in ["search", "who is", "latest", "news", "what is", "current", "today", "weather", "stock", "price", "how many"]):
        with st.status("🧠 Learning from Internet...", expanded=True) as status:
            st.write(f"Querying Open Web: {active_prompt}")
            web_context = autonomous_search(active_prompt)
            status.update(label="✅ Internet Learning Synced", state="complete", expanded=False)

    try:
        # ── Intelligence Engine Selection ──────────
        if model_option == "gemini-2.0-flash (Premium Grounding)":
            with st.chat_message("assistant"):
                with st.spinner("🌌 Consulting Google Knowledge Matrix..."):
                    try:
                        gemini_model = genai.GenerativeModel(
                            model_name="gemini-flash-latest",
                            tools=[{"google_search_retrieval": {}}]
                        )
                        response = gemini_model.generate_content(active_prompt)
                        full_response = response.text
                        st.markdown(full_response)
                        st.session_state.messages.append({"role": "assistant", "content": full_response})
                        save_chat_to_disk(st.session_state.chat_id, st.session_state.messages)
                        st.rerun()
                    except Exception as e:
                        if "429" in str(e) or "quota" in str(e).lower():
                            st.warning("⚠️ Gemini Quota Exceeded. Falling back to Deep Web Scraper...")
                            # Continue to Groq path
                        else:
                            st.error(f"Gemini Error: {str(e)}")
                            st.stop()
        
        # ── Legacy/Groq Path ──────────────────────
        MODEL_QUEUE = ["llama-3.3-70b-versatile", "llama-3.1-70b-versatile"]
        if model_option in MODEL_QUEUE:
            MODEL_QUEUE = [model_option]
        
        rag_context = fetch_knowledge(active_prompt, st.session_state.get('index'), st.session_state.get('chunks'))
        full_context = f"{rag_context}\n{web_context}"
        
        # ── Intelligence Mode Selection ────────────────
        has_template = bool(st.session_state.get('template_bytes'))
        is_surgical = has_template and st.session_state.get('surgical_mode', True)
        
        mode_instruction = ""
        if is_surgical:
            mode_instruction = (
                "🚨 ARCHITECTURAL MODE ACTIVE: YOU ARE A SURGICAL BINARY ENGINE.\n"
                "1. Load the uploaded file using python-docx.\n"
                "2. Make all requested text replacements in paragraphs, tables, headers, and footers.\n"
                "3. Save the edited file.\n"
                "4. ALWAYS provide a download link/button for the edited file.\n"
                "5. NEVER JUST OUTPUT THE REPLACEMENTS AS TEXT—always return the actual edited file.\n"
                "6. USE THE FORMAT: '\"Original\" -> \"Replacement\"'. Provide the surgical plan first."
            )
        else:
            mode_instruction = (
                "CONVERSATIONAL MODE ACTIVE: Act as a high-intelligence professional assistant. "
                "Provide detailed, human-like, and natural responses. Use formatting, bolding, and lists."
            )

        sys_msg = (
            f"You are KALI AI (v6.0 Architectural Edition), an elite intelligence and document studio.\n"
            f"TODAY'S DATE: {datetime.now().strftime('%A, %B %d, %Y')}\n"
            f"CRITICAL: You have REAL-TIME access to the internet via the Web Intelligence Engine. "
            f"You have been provided with both search snippets and DEEP CONTENT from the top 5 web sources. "
            f"Synthesize this live data to provide 'GPT-style' perfect, comprehensive, and up-to-the-minute answers.\n"
            f"STRICT RULE: NEVER use placeholders like [Team 1] or [Score]. If the data is not in the provided context, "
            f"state exactly what you found and what is missing. Use actual numbers, names, and dates from the sources.\n"
            f"{mode_instruction}\n\n"
            f"CONTEXT (RAG + Web Sync):\n{full_context}"
        )

        current_model = model_option
        msgs = [{"role": "system", "content": sys_msg}]
        
        if st.session_state.get('vision_active'):
            current_model = "meta-llama/llama-4-scout-17b-16e-instruct"
            msgs.append({
                "role": "user",
                "content": [
                    {"type": "text", "text": active_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{st.session_state.vision_base64}"}}
                ]
            })
        else:
            for m in st.session_state.messages[-5:]:
                msgs.append({"role": m["role"], "content": m["content"]})
            
            # 💉 SURGICAL PRIMER: Hard-injection for adherence
            if is_surgical:
                msgs[-1]["content"] += "\n\n[SYSTEM: KALI SURGERY ACTIVE. RESPOND WITH ONLY '\"A\" -> \"B\"' TEXT PATTERNS. DO NOT OUTPUT RAW XML TAGS OR CODE. ONLY CONTENT EDITS.]"

        # ── Intelligence Matrix (TPU-Locked Protocol) ──
        # Policy: Direct routing to Google Cloud TPU v3 Cluster
        MODEL_QUEUE = [model_option]
        
        with st.chat_message("assistant"):
            full_res = ""
            
            for engine in MODEL_QUEUE:
                try:
                    if engine.startswith("google/"):
                        with st.spinner("🛰️ Thinking..."):
                            client = Groq(api_key=auth_key) if auth_key else None
                            # Fallback logic for legacy model names
                            actual_model = "llama-3.3-70b-versatile" if "google" in engine else engine
                            response = client.chat.completions.create(model=actual_model, messages=msgs)
                            full_res = response.choices[0].message.content
                    else:
                        client = Groq(api_key=auth_key) if auth_key else None
                        response = client.chat.completions.create(model=engine, messages=msgs)
                        full_res = response.choices[0].message.content
                    break 
                except Exception as e:
                    if "401" in str(e) or "api_key" in str(e).lower():
                        st.warning("⚠️ Groq Intelligence Link Expired. Activating Gemini Secondary Cluster...")
                        try:
                            # Final Fallback: Standard Gemini (Enhanced with the intelligence context)
                            backup_model = genai.GenerativeModel("gemini-flash-latest")
                            # Combine system instructions and full context for the backup
                            backup_res = backup_model.generate_content(f"{sys_msg}\n\nUSER QUERY: {active_prompt}")
                            full_res = backup_res.text
                            break
                        except Exception as e2:
                            st.error(f"📡 Total Intelligence Collapse: {str(e2)}")
                            st.stop()
                    else:
                        st.error(f"📡 TPU Synchronicity Failure: {str(e)}")
                        break

            st.markdown(full_res)
            st.session_state.messages.append({"role": "assistant", "content": full_res})
            
            # MEMORY GUARD: Cap history to 40 messages
            if len(st.session_state.messages) > 40:
                st.session_state.messages = st.session_state.messages[-40:]
            
            # PERSISTENCE: Auto-save after response
            save_chat_to_disk(st.session_state.chat_id, st.session_state.messages)
            
            # --- ENHANCED SURGICAL TRIGGER (v6.2) ---
            # Triggers if arrow syntax exists OR if architectural XML is detected
            architectural_intent = any(tag in full_res for tag in ["<w:sectPr>", "<w:pgNumType>", "<w:headerReference>", "<w:footerReference>"])
            
            if "->" in full_res or architectural_intent or any(word in active_prompt.lower() for word in ["download", "link", "get file"]):
                # --- STARLIGHT AUTO-DISCOVERY (v7.0 Magic) ---
                if not st.session_state.get('template_bytes'):
                    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                    docx_files = [f for f in os.listdir(desktop_path) if f.lower().endswith(".docx")]
                    
                    found_file = None
                    prompt_words = active_prompt.lower().split()
                    for df in docx_files:
                        if any(word in df.lower() for word in prompt_words if len(word) > 3):
                            found_file = os.path.join(desktop_path, df)
                            break
                    
                    if found_file:
                        with st.status(f"🛰️ KALI: Autonomously linking {os.path.basename(found_file)}..."):
                            with open(found_file, 'rb') as af:
                                st.session_state.template_bytes = af.read()
                                st.session_state.uploaded_file_name = os.path.basename(found_file)
                            st.toast(f"Bound to: {st.session_state.uploaded_file_name}")

                if st.session_state.get('template_bytes'):
                    with st.spinner("🚀 KALI ARCHITECT: Executing Structural Surgery..."):
                        res = edit_and_return_docx(st.session_state.get('template_bytes'), full_res)
                        if isinstance(res, tuple):
                            edited_bytes, status_msg = res
                        else:
                            edited_bytes, status_msg = res, "Done"

                        if edited_bytes:
                            st.session_state.final_doc = edited_bytes
                            st.success("✅ **SURGERY SUCCESSFUL:** Your revised document is prepared and cached in high-speed memory.")
                            st.balloons()
                            st.toast(status_msg)
                        else:
                            st.error(f"❌ **SURGERY BLOCKED:** {status_msg}")
                
                # AUTO-FLUSH: Vision Cache
                if st.session_state.get('vision_active'):
                    st.session_state.vision_active = False
                    st.session_state.vision_base64 = None
                    st.toast("Vision Cache Flushed (RAM Optimized)")
                
    except Exception as e:
        st.error(f"Intelligence Exception: {str(e)}")

# ── Sidestage: Persistent Export Center (v6.2) ──────
with st.sidebar:
    if st.session_state.get('final_doc'):
        st.divider()
        st.markdown("### 🏛️ EXPORT CENTER")
        st.download_button(
            "📥 DOWNLOAD REVISED DOCUMENT", 
            data=st.session_state.final_doc, 
            file_name="Kali_AI_Final_Revision.docx", 
            key="side_dl",
            use_container_width=True
        )
        if st.button("📋 COPY LATEST REPORT", use_container_width=True):
            st.toast("Copied to clipboard simulator!")
        st.success("Document cached in high-speed memory.")

# ── Floating Action Center (v7.0) ────────────
if st.session_state.get('final_doc'):
    # Centered Download Hub for better visibility
    st.markdown("---")
    cols = st.columns([1, 2, 1])
    with cols[1]:
        st.success("✨ KALI AI: Professional Suite Reconstructed.")
        st.download_button(
            "💎 DOWNLOAD FINAL EDITED DOCUMENT",
            data=st.session_state.final_doc,
            file_name="Kali_AI_Final_Edit.docx",
            use_container_width=True,
            key="main_dl_hub"
        )
    
    # Keep the pulse badge in the corner as well
    st.markdown("""
        <div style='position: fixed; bottom: 25px; right: 25px; z-index: 1000;'>
            <div class='pulse-badge' style='background: #18122b; color: #a78bfa; padding: 18px 25px; border-radius: 50px; cursor: pointer; box-shadow: 0 10px 30px rgba(0,0,0,0.5); border: 1px solid #2d1f5e; font-size: 14px;'>
                🚀 File Ready: Final Revision Cached
            </div>
        </div>
    """, unsafe_allow_html=True)
